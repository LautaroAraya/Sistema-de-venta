from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Crear documento
doc = Document()

# Establecer márgenes
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Título principal
title = doc.add_heading('MANUAL DE USO', 0)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
title_format = title.runs[0]
title_format.font.size = Pt(28)
title_format.font.color.rgb = RGBColor(37, 99, 235)

subtitle = doc.add_heading('SISTEMA DE VENTAS', level=2)
subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Tabla de contenido
doc.add_heading('TABLA DE CONTENIDO', level=1)
toc_items = [
    '1. Instalación',
    '2. Primer Uso',
    '3. Guía de Uso',
    '4. Gestión de Productos',
    '5. Gestión de Proveedores',
    '6. Reportes de Ventas',
    '7. Gestión de Usuarios',
    '8. Generación de Facturas PDF',
    '9. Respaldo de Datos',
    '10. Solución de Problemas'
]
for item in toc_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# 1. INSTALACIÓN
doc.add_heading('1. INSTALACIÓN', level=1)

doc.add_heading('Para Usuarios Finales (Instalador)', level=2)
steps = [
    'Ejecutar SistemaVentas_Setup.exe',
    'Seguir las instrucciones del instalador',
    'El sistema se instalará en C:\\Program Files\\Sistema de Ventas',
    'Se creará un acceso directo en el escritorio y menú inicio',
    'Ejecutar el sistema desde el acceso directo'
]
for i, step in enumerate(steps, 1):
    doc.add_paragraph(step, style='List Number')

doc.add_heading('Para Desarrollo', level=2)
dev_steps = [
    'Tener instalado Python 3.8 o superior',
    'Abrir terminal en la carpeta del proyecto',
    'Ejecutar: pip install -r requirements.txt',
    'Ejecutar: python inicializar_datos.py (primera vez)',
    'Ejecutar: python main.py',
    'O simplemente ejecutar run.bat'
]
for step in dev_steps:
    doc.add_paragraph(step, style='List Number')

doc.add_page_break()

# 2. PRIMER USO
doc.add_heading('2. PRIMER USO', level=1)

doc.add_heading('Credenciales por Defecto', level=2)

# Tabla de credenciales
table = doc.add_table(rows=3, cols=2)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Tipo de Usuario'
hdr_cells[1].text = 'Credenciales'

row_cells = table.rows[1].cells
row_cells[0].text = 'Administrador'
row_cells[1].text = 'Usuario: admin\nContraseña: admin123'

row_cells = table.rows[2].cells
row_cells[0].text = 'Empleado (ejemplo)'
row_cells[1].text = 'Usuario: empleado1\nContraseña: empleado123'

warning = doc.add_paragraph()
warning_run = warning.add_run('⚠️  IMPORTANTE: ')
warning_run.font.bold = True
warning.add_run('Cambiar las contraseñas después del primer login.')

doc.add_page_break()

# 3. GUÍA DE USO
doc.add_heading('3. GUÍA DE USO', level=1)

doc.add_heading('Login', level=2)
login_steps = [
    'Ingresar usuario y contraseña',
    'Presionar "Iniciar Sesión" o Enter',
    'El sistema validará las credenciales'
]
for step in login_steps:
    doc.add_paragraph(step, style='List Number')

doc.add_heading('Interfaz Principal', level=2)
doc.add_paragraph('Menú Lateral (disponible para todos)', style='List Bullet')
menu_items = [
    'Nueva Venta: Módulo de facturación',
    'Productos: Gestión del inventario',
    'Proveedores: Gestión de proveedores',
    'Reportes: Historial y estadísticas de ventas',
    'Usuarios: Gestión de usuarios (SOLO ADMIN)'
]
for item in menu_items:
    doc.add_paragraph(item, style='List Bullet 2')

doc.add_paragraph('Barra Superior', style='List Bullet')
bar_items = [
    'Muestra el usuario actual y su rol',
    'Botón "Cerrar Sesión"'
]
for item in bar_items:
    doc.add_paragraph(item, style='List Bullet 2')

doc.add_page_break()

# 4. NUEVA VENTA
doc.add_heading('4. NUEVA VENTA / FACTURACIÓN', level=1)

doc.add_heading('Proceso de Venta', level=2)

sale_steps = [
    ('Datos del Cliente (Opcional)', 
     'Nombre del cliente\nDocumento de identidad\nSi no se ingresa, se asigna "Cliente Genérico"'),
    ('Buscar Producto',
     'Escribir código o nombre del producto\nLos resultados aparecen automáticamente\nClick en el producto deseado'),
    ('Configurar Item',
     'Verificar precio y stock\nIngresar cantidad deseada\nIngresar descuento % (si aplica)\nClick en "Agregar a la Venta"'),
    ('Agregar Más Productos',
     'Repetir pasos anteriores para cada producto\nLos items se van agregando a la tabla\nEl total se calcula automáticamente'),
    ('Revisar Venta',
     'Verificar todos los items en la tabla\nVerificar totales (Subtotal, Descuento, Total)\nPara eliminar un item: seleccionarlo y click en "Eliminar Item Seleccionado"'),
    ('Procesar Venta',
     'Click en "Procesar Venta"\nConfirmar la operación\nSe generará un número de factura automáticamente\nOpción de generar PDF de la factura'),
    ('Cancelar Venta',
     'Click en "Cancelar / Nueva Venta"\nConfirmar cancelación\nSe limpiará toda la información')
]

for step_title, step_content in sale_steps:
    p = doc.add_paragraph(step_title, style='List Number')
    p_run = p.runs[0]
    p_run.font.bold = True
    for line in step_content.split('\n'):
        doc.add_paragraph(line, style='List Bullet 2')

doc.add_page_break()

# 5. GESTIÓN DE PRODUCTOS
doc.add_heading('5. GESTIÓN DE PRODUCTOS', level=1)

doc.add_heading('Listar Productos', level=2)
doc.add_paragraph('Se muestran todos los productos activos')
doc.add_paragraph('Información visible: Código, Nombre, Categoría, Precio, Stock, Proveedor', style='List Bullet')

doc.add_heading('Nuevo Producto', level=2)
doc.add_paragraph('Click en "Nuevo Producto"', style='List Number')
doc.add_paragraph('Completar información:', style='List Number')
fields = [
    'Código: Único, ej: "BEB001"',
    'Nombre: Descriptivo',
    'Descripción: Opcional',
    'Categoría: Seleccionar de la lista',
    'Precio: Precio de venta',
    'Stock: Cantidad inicial',
    'Proveedor: Seleccionar de la lista'
]
for field in fields:
    doc.add_paragraph(field, style='List Bullet 2')
doc.add_paragraph('Click en "Guardar"', style='List Number')

doc.add_heading('Editar Producto', level=2)
edit_steps = [
    'Seleccionar producto en la tabla',
    'Click en "Editar"',
    'Modificar información',
    'Click en "Guardar"'
]
for step in edit_steps:
    doc.add_paragraph(step, style='List Number')

doc.add_heading('Eliminar Producto', level=2)
del_steps = [
    'Seleccionar producto en la tabla',
    'Click en "Eliminar"',
    'Confirmar eliminación',
    'El producto se desactiva (no se borra)'
]
for step in del_steps:
    doc.add_paragraph(step, style='List Number')

note = doc.add_paragraph()
note_run = note.add_run('Nota: ')
note_run.font.bold = True
note.add_run('El stock se actualiza automáticamente con cada venta.')

doc.add_page_break()

# 6. GESTIÓN DE PROVEEDORES
doc.add_heading('6. GESTIÓN DE PROVEEDORES', level=1)

doc.add_heading('Nuevo Proveedor', level=2)
doc.add_paragraph('Click en "Nuevo Proveedor"', style='List Number')
doc.add_paragraph('Completar información:', style='List Number')
prov_fields = [
    'Nombre: Nombre de la empresa (obligatorio)',
    'Contacto: Nombre del contacto',
    'Teléfono: Número de contacto',
    'Email: Correo electrónico',
    'Dirección: Dirección física'
]
for field in prov_fields:
    doc.add_paragraph(field, style='List Bullet 2')
doc.add_paragraph('Click en "Guardar"', style='List Number')

doc.add_paragraph('Editar/Eliminar: Similar al proceso de productos')

doc.add_page_break()

# 7. REPORTES DE VENTAS
doc.add_heading('7. REPORTES DE VENTAS', level=1)

doc.add_heading('Filtros Disponibles', level=2)
doc.add_paragraph('Por Fecha: Desde - Hasta (formato: YYYY-MM-DD)', style='List Bullet')
doc.add_paragraph('Filtros Rápidos:', style='List Bullet')
quick_filters = [
    'Hoy',
    'Esta Semana',
    'Este Mes',
    'Limpiar (ver todas)'
]
for f in quick_filters:
    doc.add_paragraph(f, style='List Bullet 2')

doc.add_heading('Estadísticas', level=2)
doc.add_paragraph('Total de ventas (cantidad)', style='List Bullet')
doc.add_paragraph('Total ingresos ($)', style='List Bullet')
doc.add_paragraph('Promedio por venta ($)', style='List Bullet')

doc.add_heading('Ver Detalles de Venta', level=2)
detail_steps = [
    'Seleccionar venta en la tabla',
    'Click en "Ver Detalles de Venta"',
    'Se muestra información de la factura, cliente, vendedor, productos vendidos y totales'
]
for step in detail_steps:
    doc.add_paragraph(step, style='List Number')

doc.add_page_break()

# 8. GESTIÓN DE USUARIOS
doc.add_heading('8. GESTIÓN DE USUARIOS (SOLO ADMIN)', level=1)

doc.add_heading('Crear Usuario', level=2)
user_steps = [
    'Click en "Nuevo Usuario"',
    'Completar información:'
]
doc.add_paragraph(user_steps[0], style='List Number')
doc.add_paragraph(user_steps[1], style='List Number')
user_fields = [
    'Usuario: Nombre de usuario único',
    'Nombre Completo: Nombre real del usuario',
    'Contraseña: Mínimo 6 caracteres',
    'Confirmar: Repetir contraseña',
    'Rol: Admin o Empleado'
]
for field in user_fields:
    doc.add_paragraph(field, style='List Bullet 2')
doc.add_paragraph('Click en "Guardar"', style='List Number')

doc.add_heading('Cambiar Contraseña', level=2)
pwd_steps = [
    'Seleccionar usuario',
    'Click en "Cambiar Contraseña"',
    'Ingresar nueva contraseña',
    'Confirmar contraseña',
    'Click en "Guardar"'
]
for step in pwd_steps:
    doc.add_paragraph(step, style='List Number')

doc.add_heading('Diferencias entre Roles', level=2)
doc.add_paragraph('Admin: Acceso completo al sistema', style='List Bullet')
doc.add_paragraph('Empleado: No puede gestionar productos, proveedores ni usuarios. Solo puede hacer ventas y ver reportes', style='List Bullet')

doc.add_page_break()

# 9. FACTURAS PDF
doc.add_heading('9. GENERACIÓN DE FACTURAS PDF', level=1)

doc.add_paragraph('Las facturas se generan automáticamente en formato PDF después de procesar una venta.')

doc.add_heading('Ubicación', level=2)
doc.add_paragraph('Desarrollo: carpeta_proyecto/reports/', style='List Bullet')
doc.add_paragraph('Instalado: C:\\Program Files\\Sistema de Ventas\\reports\\', style='List Bullet')

doc.add_heading('Nombre', level=2)
doc.add_paragraph('FAC-YYYYMMDD-XXXX.pdf')

doc.add_heading('Contenido de la Factura', level=2)
pdf_items = [
    'Número de factura',
    'Fecha y hora',
    'Datos del vendedor',
    'Datos del cliente',
    'Detalle de productos',
    'Totales'
]
for item in pdf_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# 10. RESPALDO DE DATOS
doc.add_heading('10. RESPALDO DE DATOS', level=1)

doc.add_heading('Base de Datos', level=2)
doc.add_paragraph('Desarrollo: carpeta_proyecto/database/sistema_ventas.db', style='List Bullet')
doc.add_paragraph('Instalado: C:\\Program Files\\Sistema de Ventas\\database\\sistema_ventas.db', style='List Bullet')

doc.add_heading('Hacer Respaldo', level=2)
backup_steps = [
    'Copiar el archivo sistema_ventas.db',
    'Guardarlo en un lugar seguro',
    'Incluir la fecha en el nombre: sistema_ventas_2026-01-15.db'
]
for step in backup_steps:
    doc.add_paragraph(step, style='List Number')

doc.add_heading('Restaurar Respaldo', level=2)
restore_steps = [
    'Cerrar el sistema',
    'Reemplazar sistema_ventas.db con el respaldo',
    'Reiniciar el sistema'
]
for step in restore_steps:
    doc.add_paragraph(step, style='List Number')

recommendation = doc.add_paragraph()
rec_run = recommendation.add_run('RECOMENDACIÓN: ')
rec_run.font.bold = True
recommendation.add_run('Hacer respaldos diarios de la base de datos.')

doc.add_page_break()

# 11. SOLUCIÓN DE PROBLEMAS
doc.add_heading('11. SOLUCIÓN DE PROBLEMAS', level=1)

problems = [
    ('"Usuario o contraseña incorrectos"', 'Verificar mayúsculas/minúsculas\nSi olvidó la contraseña de admin, contactar soporte'),
    ('"Stock insuficiente"', 'Verificar stock disponible del producto\nActualizar stock en Gestión de Productos'),
    ('"Error al crear venta"', 'Verificar que todos los productos tengan stock\nReiniciar el sistema\nVerificar integridad de la base de datos'),
    ('La aplicación no inicia', 'Verificar que Python esté instalado (desarrollo)\nReinstalar la aplicación (instalador)\nRevisar logs de errores'),
    ('Facturas no se generan', 'Verificar permisos de escritura en carpeta reports\nVerificar que ReportLab esté instalado')
]

for problem, solution in problems:
    doc.add_heading(problem, level=2)
    for line in solution.split('\n'):
        doc.add_paragraph(line, style='List Bullet')

doc.add_page_break()

# Pie de página
doc.add_heading('INFORMACIÓN ADICIONAL', level=1)

doc.add_heading('Atajos de Teclado', level=2)
doc.add_paragraph('Enter: Aceptar/Continuar', style='List Bullet')
doc.add_paragraph('Esc: Cancelar (en diálogos)', style='List Bullet')
doc.add_paragraph('Tab: Navegar entre campos', style='List Bullet')

doc.add_heading('Consejos y Mejores Prácticas', level=2)
tips = [
    ('Seguridad', 'Cambiar contraseñas por defecto\nUsar contraseñas seguras (min. 8 caracteres)\nCerrar sesión al terminar'),
    ('Inventario', 'Actualizar stock regularmente\nRevisar productos con stock bajo\nMantener códigos de producto consistentes'),
    ('Ventas', 'Siempre verificar totales antes de procesar\nGuardar facturas PDF importantes\nRevisar reportes diariamente'),
    ('Respaldos', 'Hacer respaldos diarios de la base de datos\nGuardar respaldos en múltiples ubicaciones\nProbar restaurar respaldos periódicamente')
]

for tip_title, tip_content in tips:
    doc.add_heading(tip_title, level=3)
    for line in tip_content.split('\n'):
        doc.add_paragraph(line, style='List Bullet')

# Pie de página final
doc.add_page_break()
footer_para = doc.add_paragraph()
footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
footer_para.add_run('Versión del Manual: 1.0\n')
footer_para.add_run('Fecha: Enero 2026\n')
footer_para.add_run('Sistema: Sistema de Ventas v1.0')

# Guardar documento
doc.save('MANUAL_USO.docx')
print("✅ Manual de uso generado exitosamente: MANUAL_USO.docx")
